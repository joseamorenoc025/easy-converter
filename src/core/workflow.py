import os
import shutil
from enum import Enum
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt

class RuleType(Enum):
    RENAME = "rename"
    MOVE = "move"
    COPY = "copy"
    NOTIFY = "notify"
    SANITIZE = "sanitize"

@dataclass
class WorkflowRule:
    type: RuleType
    params: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self):
        return {
            "type": self.type.value,
            "params": self.params
        }

    @classmethod
    def from_dict(cls, data):
        return cls(type=RuleType(data["type"]), params=data.get("params", {}))

@dataclass
class WorkflowProfile:
    name: str
    rules: List[WorkflowRule] = field(default_factory=list)
    watch_path: Optional[str] = None
    is_active: bool = False

    def to_dict(self):
        return {
            "name": self.name,
            "rules": [rule.to_dict() for rule in self.rules],
            "watch_path": self.watch_path,
            "is_active": self.is_active
        }

    @classmethod
    def from_dict(cls, data):
        return cls(
            name=data["name"],
            rules=[WorkflowRule.from_dict(r) for r in data.get("rules", [])],
            watch_path=data.get("watch_path"),
            is_active=data.get("is_active", False)
        )

class WorkflowManager:
    def __init__(self, config_manager):
        self.config_manager = config_manager
        self.profiles: List[WorkflowProfile] = self._load_profiles()

    def _load_profiles(self) -> List[WorkflowProfile]:
        data = self.config_manager.get("workflow_profiles", [])
        return [WorkflowProfile.from_dict(p) for p in data]

    def save_profiles(self):
        self.config_manager.set("workflow_profiles", [p.to_dict() for p in self.profiles])

    def add_profile(self, profile: WorkflowProfile):
        self.profiles.append(profile)
        self.save_profiles()

    def delete_profile(self, name: str):
        self.profiles = [p for p in self.profiles if p.name != name]
        self.save_profiles()

    def apply_workflow(self, profile_name: str, file_path: Path) -> Path:
        profile = next((p for p in self.profiles if p.name == profile_name), None)
        if not profile:
            return file_path

        current_path = file_path
        for rule in profile.rules:
            try:
                current_path = self._execute_rule(rule, current_path)
            except Exception as e:
                print(f"Error aplicando regla {rule.type}: {e}")
        
        return current_path

    def _execute_rule(self, rule: WorkflowRule, file_path: Path) -> Path:
        if rule.type == RuleType.RENAME:
            pattern = rule.params.get("pattern", "{name}_{date}")
            date_str = datetime.now().strftime("%Y%m%d_%H%M%S")
            new_name = pattern.replace("{name}", file_path.stem).replace("{date}", date_str) + file_path.suffix
            new_path = file_path.parent / new_name
            os.rename(file_path, new_path)
            return new_path

        elif rule.type == RuleType.MOVE:
            dest_dir = Path(rule.params.get("destination"))
            dest_dir.mkdir(parents=True, exist_ok=True)
            dest_path = dest_dir / file_path.name
            shutil.move(str(file_path), str(dest_path))
            return dest_path

        elif rule.type == RuleType.COPY:
            dest_dir = Path(rule.params.get("destination"))
            dest_dir.mkdir(parents=True, exist_ok=True)
            dest_path = dest_dir / file_path.name
            shutil.copy2(str(file_path), str(dest_path))
            return file_path # Retornamos el original para seguir aplicando reglas sobre él

        elif rule.type == RuleType.SANITIZE:
            if file_path.suffix.lower() != '.docx':
                return file_path
            
            try:
                doc = Document(str(file_path))
                font_name = rule.params.get("font", "Arial")
                font_size = rule.params.get("size", 11)
                
                for p in doc.paragraphs:
                    # Unir líneas cortadas por saltos manuales
                    if "\n" in p.text:
                        p.text = p.text.replace("\n", " ")
                    
                    # Unificar formato de fuente
                    for run in p.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                
                # Eliminar párrafos vacíos (colección separada para evitar corrupción del árbol XML)
                empty_paras = [p for p in doc.paragraphs if not p.text.strip()]
                for p in empty_paras:
                    p._element.getparent().remove(p._element)
                
                doc.save(str(file_path))
            except Exception as e:
                print(f"Error sanitizando documento: {e}")
            
            return file_path

        return file_path
