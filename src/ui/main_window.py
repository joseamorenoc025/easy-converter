import customtkinter
from pathlib import Path
from tkinterdnd2 import TkinterDnD
from PIL import Image
import fitz
from core.converter import EasyConverter
from core.controller import AppController
from core.converter_adapter import ConverterAdapter
from core.queue_adapter import QueueAdapter
from core.config_adapter import ConfigAdapter
from core.workflow_adapter import WorkflowAdapter
from core.platform_adapter import PlatformAdapter
from core.view_adapter import ViewAdapter
from core.error_handler import ErrorHandler
from core.progress import ConversionProgress
from core.queue_manager import QueueItem
from core.file_manager import PathManager
from core.workflow import WorkflowProfile
from core.watcher import SmartWatcher
from ui.workflow_panel import WorkflowPanel
from ui.pdf_operations import MergePanel, SplitPanel, CompressPanel, SanitizePanel, DecryptPanel, BatchPanel, EncryptPanel
from ui.themes import ThemeManager
from ui.components import FileDropZone, ModeToggle, ProgressCard
from ui.notifications import NotificationManager
from utils.context_menu import add_context_menu
from utils.pdf_tools import extract_text_with_ocr
from utils.security import is_safe_path, validate_file_magic
from tkinter import messagebox

customtkinter.set_default_color_theme("blue")


class SidebarSection(customtkinter.CTkFrame):
    """Sección visual del sidebar con título y separador."""
    def __init__(self, master, title, **kwargs):
        from ui.themes import ThemeManager
        super().__init__(master, fg_color="transparent", **kwargs)

        sep = customtkinter.CTkFrame(self, height=1, fg_color=ThemeManager.get_color("border"))
        sep.pack(fill="x", padx=15, pady=(10, 4))

        lbl = customtkinter.CTkLabel(self, text=title,
                                     font=customtkinter.CTkFont(size=11, weight="bold"),
                                     text_color=ThemeManager.get_color("text_secondary"),
                                     anchor="w")
        lbl.pack(padx=20, pady=(0, 4), anchor="w")


class App(customtkinter.CTk, TkinterDnD.DnDWrapper):
    def __init__(self):
        super().__init__()
        self.TkdndVersion = TkinterDnD._require(self)

        self.title("Easy Converter v2.2")
        self.geometry("900x700")
        self.resizable(True, True)
        self.minsize(750, 600)
        self._set_app_icon()

        self.config_adapter = ConfigAdapter()
        ThemeManager.apply(self.config_adapter.get("theme", "dark"))
        self.path_manager = PathManager(self.config_adapter._config)
        self.error_handler = ErrorHandler()
        self.watcher = SmartWatcher(callback=self.on_file_detected)

        converter = ConverterAdapter()
        queue = QueueAdapter()
        queue.set_converter(converter)
        self.controller = AppController(
            converter=converter,
            queue_manager=queue,
            config_manager=self.config_adapter,
            workflow_engine=WorkflowAdapter(self.config_adapter._config),
            platform_service=PlatformAdapter(),
            view_callback=ViewAdapter(self),
        )
        self.queue_manager = self.controller._queue_manager._queue
        self.config_manager = self.config_adapter._config
        self.workflow_manager = self.controller._workflow_engine._manager
        self.notifications = NotificationManager(master=self)
        self.conversion_mode = "pdf2word"

        self.setup_ui()
        self.setup_watchers()

    def _set_app_icon(self):
        icon_path = Path(__file__).parent.parent / "assets" / "icon.ico"
        if icon_path.exists():
            try:
                self.iconbitmap(str(icon_path))
            except Exception:
                pass

    def setup_watchers(self):
        self.watcher.start()
        for profile in self.workflow_manager.profiles:
            if profile.is_active and profile.watch_path:
                self.watcher.add_watch(profile.watch_path)

    def on_file_detected(self, file_path: Path):
        if not is_safe_path(file_path):
            self.notifications.error("Archivo no permitido", f"{file_path}")
            return

        ext = file_path.suffix.lower()
        expected_type = 'pdf' if ext == '.pdf' else ('docx' if ext in ['.docx', '.doc'] else None)
        if expected_type:
            is_valid, msg = validate_file_magic(file_path, expected_type)
            if not is_valid:
                self.notifications.error("Archivo inv\u00e1lido", msg)
                return

        for profile in self.workflow_manager.profiles:
            if profile.is_active and profile.watch_path:
                if str(file_path.parent).lower() == str(Path(profile.watch_path)).lower():
                    mode = 'pdf2word' if ext == '.pdf' else 'word2pdf'
                    self.after(0, lambda fp=file_path: self.queue_manager.add_item(fp, mode, workflow_profile=profile.name))
                    break

    def handle_watcher_change(self, profile: WorkflowProfile):
        if profile.is_active and profile.watch_path:
            self.watcher.add_watch(profile.watch_path)
        elif profile.watch_path:
            self.watcher.remove_watch(profile.watch_path)

    def _build_sidebar(self):
        self.sidebar = customtkinter.CTkFrame(self, width=210, corner_radius=0,
                                              fg_color=ThemeManager.get_sidebar_bg())
        self.sidebar.pack(side="left", fill="y")
        self.sidebar.pack_propagate(False)

        logo_frame = customtkinter.CTkFrame(self.sidebar, fg_color="transparent")
        logo_frame.pack(fill="x", pady=(18, 8), padx=18)
        customtkinter.CTkLabel(logo_frame, text="\u2699\ufe0f EasyConverter",
                               font=customtkinter.CTkFont(size=18, weight="bold")).pack(anchor="w")
        customtkinter.CTkLabel(logo_frame, text="v2.2",
                               font=customtkinter.CTkFont(size=10),
                               text_color=ThemeManager.get_color("text_secondary")).pack(anchor="w")

        sec_output = SidebarSection(self.sidebar, title="\U0001f4c2  Salida")
        sec_output.pack(fill="x")

        out_frame = customtkinter.CTkFrame(self.sidebar, fg_color="transparent")
        out_frame.pack(fill="x", padx=18, pady=(0, 6))

        self.out_mode_var = customtkinter.StringVar(value=self.config_manager.get("output_mode"))
        self.radio_same = customtkinter.CTkRadioButton(out_frame, text="Misma carpeta",
                                                       variable=self.out_mode_var, value="same_folder",
                                                       command=self.save_settings)
        self.radio_same.pack(pady=3, anchor="w")
        self.radio_sub = customtkinter.CTkRadioButton(out_frame, text="Subcarpeta 'convertidos'",
                                                      variable=self.out_mode_var, value="subfolder",
                                                      command=self.save_settings)
        self.radio_sub.pack(pady=3, anchor="w")
        self.radio_custom = customtkinter.CTkRadioButton(out_frame, text="Carpeta personalizada",
                                                         variable=self.out_mode_var, value="custom",
                                                         command=self.save_settings)
        self.radio_custom.pack(pady=3, anchor="w")

        self.btn_select_custom = customtkinter.CTkButton(out_frame, text="Seleccionar Carpeta",
                                                         command=self.select_custom_folder,
                                                         height=26, font=customtkinter.CTkFont(size=11),
                                                         fg_color=ThemeManager.get_color("primary"))
        self.btn_select_custom.pack(pady=(6, 0))

        sec_opts = SidebarSection(self.sidebar, title="\u2699\ufe0f  Opciones")
        sec_opts.pack(fill="x")

        opts_frame = customtkinter.CTkFrame(self.sidebar, fg_color="transparent")
        opts_frame.pack(fill="x", padx=18, pady=(0, 6))

        self.check_open = customtkinter.CTkCheckBox(opts_frame, text="Abrir al finalizar",
                                                    command=self.save_settings)
        if self.config_manager.get("open_folder_on_finish"):
            self.check_open.select()
        self.check_open.pack(pady=3, anchor="w")

        self.check_ocr = customtkinter.CTkCheckBox(opts_frame, text="OCR (lento)",
                                                   command=self.save_settings)
        if self.config_manager.get("use_ocr"):
            self.check_ocr.select()
        self.check_ocr.pack(pady=3, anchor="w")

        sec_theme = SidebarSection(self.sidebar, title="\U0001f3a8  Tema")
        sec_theme.pack(fill="x")

        current_theme = self.config_manager.get("theme", "dark")
        self.theme_map = {ThemeManager.get_theme_label(k): k for k in ThemeManager.get_theme_names()}
        self.theme_var = customtkinter.StringVar(value=ThemeManager.get_theme_label(current_theme))
        self.theme_menu = customtkinter.CTkOptionMenu(
            self.sidebar, variable=self.theme_var,
            values=list(self.theme_map.keys()),
            command=self.change_theme, width=170,
            fg_color=ThemeManager.get_color("card_bg"),
            button_color=ThemeManager.get_color("primary"))
        self.theme_menu.pack(pady=(4, 10), padx=18)

        self.btn_context_menu = customtkinter.CTkButton(
            self.sidebar, text="\U0001f4cc Men\u00fa Contextual",
            command=self.register_context_menu, height=32,
            fg_color=ThemeManager.get_color("card_bg"),
            hover_color=ThemeManager.get_color("card_hover"),
            font=customtkinter.CTkFont(size=11))
        self.btn_context_menu.pack(pady=(0, 14), padx=18, side="bottom")

    def setup_ui(self):
        self._build_sidebar()

        self.main_frame = customtkinter.CTkFrame(self, fg_color="transparent")
        self.main_frame.pack(side="right", fill="both", expand=True, padx=(8, 18))

        self.tabview = customtkinter.CTkTabview(self.main_frame, corner_radius=10)
        self.tabview.pack(fill="both", expand=True, pady=(8, 0))
        self.tabview.add("\U0001f4c4  Conversi\u00f3n")
        self.tabview.add("\U0001f6e0\ufe0f  Herramientas")
        self.tabview.add("\U0001f504  Flujos")
        self.tabview.add("\U0001f4e6  Lote")

        self._build_conversion_tab()
        self._build_tools_tab()
        self._build_flows_tab()
        self._build_batch_tab()
        self._build_footer()

        self.setup_tooltips()

    def _build_conversion_tab(self):
        tab = self.tabview.tab("\U0001f4c4  Conversi\u00f3n")
        content = customtkinter.CTkFrame(tab, fg_color="transparent")
        content.pack(fill="both", expand=True)

        left = customtkinter.CTkFrame(content, fg_color="transparent")
        left.pack(side="left", fill="both", expand=True, padx=(0, 8))

        self.mode_frame = ModeToggle(left, on_change=self.set_mode)
        self.mode_frame.pack(pady=(10, 6))

        self.drop_zone = FileDropZone(left, on_drop=self.process_selected_file)
        self.drop_zone.pack(pady=8, fill="x")

        queue_header = customtkinter.CTkFrame(left, fg_color="transparent")
        queue_header.pack(fill="x", pady=(8, 2))
        customtkinter.CTkLabel(queue_header, text="Cola de conversi\u00f3n",
                               font=customtkinter.CTkFont(size=12, weight="bold"),
                               text_color=ThemeManager.get_color("text_secondary")).pack(side="left")
        self.queue_count_label = customtkinter.CTkLabel(queue_header, text="0 archivos",
                                                        font=customtkinter.CTkFont(size=11),
                                                        text_color=ThemeManager.get_color("text_secondary"))
        self.queue_count_label.pack(side="right")

        self.scroll_frame = customtkinter.CTkScrollableFrame(left, height=280,
                                                             fg_color=ThemeManager.get_color("surface"))
        self.scroll_frame.pack(fill="both", expand=True, pady=(0, 4))
        self.item_widgets = {}

        self.preview_frame = customtkinter.CTkFrame(content, width=200, corner_radius=10,
                                                    fg_color=ThemeManager.get_color("card_bg"))
        self.preview_frame.pack(side="right", fill="y", pady=8)
        self.preview_frame.pack_propagate(False)

        customtkinter.CTkLabel(self.preview_frame, text="\U0001f441\ufe0f",
                               font=customtkinter.CTkFont(size=22)).pack(pady=(10, 2))
        customtkinter.CTkLabel(self.preview_frame, text="Vista Previa",
                               font=customtkinter.CTkFont(size=11, weight="bold")).pack()

        self.preview_image_label = customtkinter.CTkLabel(self.preview_frame, text="Sin selecci\u00f3n",
                                                          text_color=ThemeManager.get_color("text_secondary"),
                                                          font=customtkinter.CTkFont(size=11))
        self.preview_image_label.pack(expand=True, padx=10, pady=10)

    def _build_tools_tab(self):
        tab = self.tabview.tab("\U0001f6e0\ufe0f  Herramientas")
        self.tools_tabview = customtkinter.CTkTabview(tab, corner_radius=8)
        self.tools_tabview.pack(fill="both", expand=True)

        tabs = [
            ("\U0001f500 Combinar", MergePanel),
            ("\u2702\ufe0f Dividir", SplitPanel),
            ("\U0001f4e6 Comprimir", CompressPanel),
            ("\U0001f6e1\ufe0f Sanitizar", SanitizePanel),
            ("\U0001f513 Descifrar", DecryptPanel),
            ("\U0001f512 Cifrar", EncryptPanel),
        ]
        for label, PanelClass in tabs:
            self.tools_tabview.add(label)
            panel = PanelClass(self.tools_tabview.tab(label))
            panel.pack(fill="both", expand=True)

    def _build_flows_tab(self):
        tab = self.tabview.tab("\U0001f504  Flujos")
        self.workflow_panel = WorkflowPanel(tab, self.workflow_manager,
                                           on_watcher_change=self.handle_watcher_change)
        self.workflow_panel.pack(fill="both", expand=True)

    def _build_batch_tab(self):
        tab = self.tabview.tab("\U0001f4e6  Lote")
        self.batch_panel = BatchPanel(tab, on_batch_convert=self._batch_convert)
        self.batch_panel.pack(fill="both", expand=True)

    def _build_footer(self):
        self.footer = customtkinter.CTkFrame(self.main_frame, fg_color="transparent", height=40)
        self.footer.pack(fill="x", pady=(4, 0))
        self.footer.pack_propagate(False)

        self.clear_button = customtkinter.CTkButton(
            self.footer, text="\U0001f5d1\ufe0f Limpiar completados",
            command=self.queue_manager.clear_completed, width=160, height=28,
            fg_color=ThemeManager.get_color("card_bg"),
            hover_color=ThemeManager.get_color("card_hover"),
            font=customtkinter.CTkFont(size=11))
        self.clear_button.pack(side="left", pady=4)

        self.status_global = customtkinter.CTkLabel(self.footer, text="\u2714 Listo",
                                                    font=customtkinter.CTkFont(size=11),
                                                    text_color=ThemeManager.get_color("success"))
        self.status_global.pack(side="right", pady=4)

    def setup_tooltips(self):
        from ui.components import ToolTip
        ToolTip(self.btn_pdf2word if hasattr(self, 'btn_pdf2word') else self.mode_frame.btn_pdf2word,
                "Convertir archivos PDF a formato Word (DOCX)")
        ToolTip(self.mode_frame.btn_word2pdf,
                "Convertir archivos Word (DOCX) a formato PDF")
        ToolTip(self.check_open, "Abrir la carpeta contenedora al finalizar la conversi\u00f3n")
        ToolTip(self.check_ocr, "Extraer texto de im\u00e1genes en PDF usando OCR (requiere Tesseract)")
        ToolTip(self.clear_button, "Eliminar de la lista los elementos completados")
        ToolTip(self.btn_context_menu, "Agregar opciones 'Convertir a Word/PDF' al men\u00fa contextual de Windows")

    def save_settings(self):
        self.config_manager.set("output_mode", self.out_mode_var.get())
        self.config_manager.set("open_folder_on_finish", self.check_open.get())
        self.config_manager.set("use_ocr", self.check_ocr.get())

    def change_theme(self, display_name):
        theme_key = self.theme_map.get(display_name)
        if not theme_key:
            return
        ThemeManager.apply(theme_key)
        self.config_manager.set("theme", theme_key)

        # Guardar estado
        queue_items = self.queue_manager.get_all_items()
        out_mode = self.out_mode_var.get()
        ocr_val = self.check_ocr.get()
        open_val = self.check_open.get()
        conv_mode = self.conversion_mode

        # Destruir y reconstruir UI
        self.sidebar.destroy()
        self.main_frame.destroy()
        self.footer.destroy()

        self._build_sidebar()

        self.main_frame = customtkinter.CTkFrame(self, fg_color="transparent")
        self.main_frame.pack(side="right", fill="both", expand=True, padx=(8, 18))

        self.tabview = customtkinter.CTkTabview(self.main_frame, corner_radius=10)
        self.tabview.pack(fill="both", expand=True, pady=(8, 0))
        self.tabview.add("\U0001f4c4  Conversi\u00f3n")
        self.tabview.add("\U0001f6e0\ufe0f  Herramientas")
        self.tabview.add("\U0001f504  Flujos")
        self.tabview.add("\U0001f4e6  Lote")

        self._build_conversion_tab()
        self._build_tools_tab()
        self._build_flows_tab()
        self._build_batch_tab()
        self._build_footer()

        # Restaurar estado
        self.out_mode_var.set(out_mode)
        if ocr_val:
            self.check_ocr.select()
        if open_val:
            self.check_open.select()
        self.conversion_mode = conv_mode

        # Restaurar cola visual
        self.item_widgets = {}
        self._render_queue()
        self.setup_tooltips()

    def select_custom_folder(self):
        folder = customtkinter.filedialog.askdirectory()
        if folder:
            self.config_manager.set("custom_path", folder)
            self.out_mode_var.set("custom")
            self.save_settings()

    def _batch_convert(self, folder_path, mode, recursive):
        return self.controller.batch_convert_folder(folder_path, mode, recursive)

    def select_files_dialog(self):
        file_paths = customtkinter.filedialog.askopenfilenames(
            title="Seleccionar archivos",
            filetypes=[("Archivos compatibles", "*.pdf *.docx *.doc"), ("PDF", "*.pdf"), ("Word", "*.docx *.doc")]
        )
        for path in file_paths:
            self.process_selected_file(path)

    def process_selected_file(self, file_path):
        path = Path(file_path)

        if not is_safe_path(path):
            self.notifications.error("Archivo no permitido", f"{path}")
            return

        ext = path.suffix.lower()
        expected_type = 'pdf' if ext == '.pdf' else ('docx' if ext in ['.docx', '.doc'] else None)
        if expected_type:
            is_valid, msg = validate_file_magic(path, expected_type)
            if not is_valid:
                self.notifications.error("Archivo inv\u00e1lido", msg)
                return

        if ext == '.pdf':
            self.update_preview(path)

        mode = None
        if ext == '.pdf':
            mode = 'pdf2word'
        elif ext in ['.docx', '.doc']:
            mode = 'word2pdf'

        if mode:
            success, message, _ = self.controller.queue_conversion(path, mode)
            if not success:
                self.notifications.error("Error en cola", f"No se pudo a\u00f1adir: {message}")

    def update_preview(self, pdf_path):
        try:
            doc = fitz.open(str(pdf_path))
            page = doc.load_page(0)
            pix = page.get_pixmap(matrix=fitz.Matrix(0.5, 0.5))
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            preview_width = 180
            aspect = img.height / img.width
            preview_height = int(preview_width * aspect)

            ctk_img = customtkinter.CTkImage(light_image=img, dark_image=img,
                                              size=(preview_width, preview_height))
            self.preview_image_label.configure(image=ctk_img, text="")
            doc.close()
        except Exception as e:
            self.preview_image_label.configure(image=None, text="Error en preview")
            print(f"Error generando preview: {e}")

    def set_mode(self, mode):
        self.conversion_mode = mode

    def process_item(self, item: QueueItem):
        output_path = self.path_manager.resolve_output_path(item.file_path, item.mode)

        def on_update(percent, message):
            item.progress = percent
            item.message = message
            self.update_queue_ui()

        tracker = ConversionProgress(on_update=on_update)
        try:
            if item.mode == "pdf2word":
                success, result = EasyConverter.pdf_to_docx(item.file_path, output_path=output_path, progress_tracker=tracker)
                if success and item.use_ocr and result:
                    item.message = "Ejecutando OCR..."
                    self.update_queue_ui()
                    try:
                        from docx import Document
                        ocr_text = extract_text_with_ocr(str(item.file_path), lang="spa")
                        if ocr_text.strip():
                            doc = Document(result)
                            doc.add_paragraph("")
                            p = doc.add_paragraph("=== TEXTO EXTRA\u00cdDO POR OCR ===")
                            for run in p.runs:
                                run.bold = True
                            for line in ocr_text.split("\n"):
                                if line.strip():
                                    doc.add_paragraph(line.strip())
                            doc.save(result)
                    except Exception as ocr_err:
                        print(f"Error en OCR: {ocr_err}")
            else:
                success, result = EasyConverter.docx_to_pdf(item.file_path, output_path=output_path, progress_tracker=tracker)

            if success:
                if item.workflow_profile:
                    item.message = "Aplicando reglas..."
                    self.update_queue_ui()
                    final_path = self.workflow_manager.apply_workflow(item.workflow_profile, Path(result))
                    item.result_path = str(final_path)
                else:
                    item.result_path = result

                self.after(0, lambda: self.notifications.success(
                    "Conversi\u00f3n exitosa",
                    f"{item.file_path.name} \u2192 {Path(item.result_path).name}"
                ))
                if self.config_manager.get("open_folder_on_finish"):
                    self.path_manager.open_in_explorer(Path(item.result_path))
        except Exception as e:
            error_msg = str(e)[:100]
            file_name = item.file_path.name
            self.after(0, lambda f=file_name, err=error_msg: self.notifications.error(
                "Error de conversi\u00f3n", f"{f}: {err}"
            ))
            raise e

    def register_context_menu(self):
        success, message = add_context_menu()
        if success:
            messagebox.showinfo("\u00c9xito", f"{message}\nAhora puedes hacer clic derecho en archivos PDF/DOCX.")
        else:
            messagebox.showerror("Error", message)

    def update_queue_ui(self):
        self.after(0, self._render_queue)

    def _render_queue(self):
        items = self.queue_manager.get_all_items()
        current_ids = [id(item) for item in items]

        for item_id in list(self.item_widgets.keys()):
            if item_id not in current_ids:
                self.item_widgets[item_id]['frame'].destroy()
                del self.item_widgets[item_id]

        alt = False
        for item in items:
            item_id = id(item)
            if item_id not in self.item_widgets:
                row_bg = ThemeManager.get_color("queue_row") if not alt else ThemeManager.get_color("queue_row_alt")
                card = ProgressCard(self.scroll_frame, item.file_path.name, mode=item.mode)
                card.configure(fg_color=row_bg)
                card.pack(fill="x", pady=3, padx=5)

                remove_btn = customtkinter.CTkButton(
                    card, text="\u2715", width=26, height=22,
                    fg_color="transparent",
                    text_color=ThemeManager.get_color("text_secondary"),
                    hover_color=ThemeManager.get_color("error"),
                    font=customtkinter.CTkFont(size=11),
                    command=lambda i=item, iid=item_id: self._remove_queue_item(i, iid))
                remove_btn.pack(side="right", padx=(0, 6), pady=6)

                self.item_widgets[item_id] = {
                    'frame': card, 'card': card,
                    'remove_btn': remove_btn,
                    '_last_msg': '', '_last_status': '', '_last_progress': -1
                }
            alt = not alt

            widgets = self.item_widgets[item_id]
            card = widgets['card']

            if item.progress != widgets['_last_progress']:
                card.progress_bar.set(item.progress / 100)
                widgets['_last_progress'] = item.progress
            if item.message != widgets['_last_msg']:
                card.status_label.configure(text=item.message)
                widgets['_last_msg'] = item.message
            if item.status != widgets['_last_status']:
                if item.status == "success":
                    card.status_label.configure(text_color=ThemeManager.get_color("success"))
                    card.open_btn.configure(state="normal")
                elif item.status == "failed":
                    card.status_label.configure(text_color=ThemeManager.get_color("error"))
                elif item.status == "running":
                    card.status_label.configure(text_color=ThemeManager.get_color("warning"))
                elif item.status == "retry_pending":
                    card.status_label.configure(text_color=ThemeManager.get_color("accent"))
                else:
                    card.status_label.configure(text_color=ThemeManager.get_color("text_secondary"))
                widgets['_last_status'] = item.status

        self.queue_count_label.configure(text=f"{len(items)} archivo{'s' if len(items) != 1 else ''}")

        running = sum(1 for i in items if i.status == "running")
        if running:
            self.status_global.configure(text=f"\u23f3 Procesando {running}...",
                                         text_color=ThemeManager.get_color("warning"))
        else:
            self.status_global.configure(text="\u2714 Listo",
                                         text_color=ThemeManager.get_color("success"))

    def _remove_queue_item(self, item, item_id):
        self.queue_manager.remove_item(item)
        if item_id in self.item_widgets:
            self.item_widgets[item_id]['frame'].destroy()
            del self.item_widgets[item_id]


if __name__ == "__main__":
    app = App()
    app.mainloop()
