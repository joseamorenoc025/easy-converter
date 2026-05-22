import unittest
import os
import sys
import tempfile
import shutil
from pathlib import Path
from unittest.mock import MagicMock

# Add src to path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'src')))

from core.workflow import WorkflowManager, WorkflowProfile, WorkflowRule, RuleType

class TestWorkflow(unittest.TestCase):
    def setUp(self):
        self.test_dir = tempfile.mkdtemp()
        self.config_manager_mock = MagicMock()
        self.config_manager_mock.get.return_value = []
        self.workflow_manager = WorkflowManager(self.config_manager_mock)
        
    def tearDown(self):
        shutil.rmtree(self.test_dir)
        
    def test_rename_rule(self):
        rule = WorkflowRule(type=RuleType.RENAME, params={"pattern": "PREFIX_{name}"})
        profile = WorkflowProfile(name="test", rules=[rule])
        self.workflow_manager.add_profile(profile)
        
        # Create a dummy file
        test_file = Path(self.test_dir) / "doc.pdf"
        test_file.touch()
        
        new_path = self.workflow_manager.apply_workflow("test", test_file)
        self.assertTrue(new_path.name.startswith("PREFIX_doc"))
        self.assertTrue(new_path.name.endswith(".pdf"))
        self.assertTrue(new_path.exists())
        self.assertFalse(test_file.exists()) # Original should be renamed
        
    def test_sanitize_skips_headings(self):
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx not installed")
            
        doc_path = Path(self.test_dir) / "test.docx"
        doc = Document()
        doc.add_heading('Header', level=1)
        doc.add_paragraph('Normal paragraph.')
        doc.save(doc_path)
        
        rule = WorkflowRule(type=RuleType.SANITIZE, params={"font": "Courier", "size": 12})
        profile = WorkflowProfile(name="test", rules=[rule])
        self.workflow_manager.add_profile(profile)
        
        new_path = self.workflow_manager.apply_workflow("test", doc_path)
        self.assertEqual(new_path, doc_path)
        
        # Read back and verify
        doc_result = Document(doc_path)
        self.assertEqual(doc_result.paragraphs[0].style.name, 'Heading 1')
        self.assertEqual(doc_result.paragraphs[0].text, 'Header')
        if len(doc_result.paragraphs[0].runs) > 0:
            self.assertNotEqual(doc_result.paragraphs[0].runs[0].font.name, 'Courier')

if __name__ == '__main__':
    unittest.main()
